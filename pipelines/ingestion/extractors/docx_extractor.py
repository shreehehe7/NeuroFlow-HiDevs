import docx
from typing import List
from . import ExtractedPage

def extract_docx(file_path: str) -> List[ExtractedPage]:
    pages: List[ExtractedPage] = []
    doc = docx.Document(file_path)
    
    current_level = None
    current_section = None
    
    # Python-docx doesn't easily map to "pages" natively, so we treat the whole document as page 1,
    # or divide it by sections if needed. For now, page_number=1 is used for everything.
    # Alternatively, a chunk can be its own pseudo-page.
    page_number = 1
    
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            try:
                # E.g., 'Heading 1' -> 'h1'
                level = f"h{para.style.name.split(' ')[1]}"
            except IndexError:
                level = "h1"
            
            current_level = level
            current_section = para.text.strip()
            
            if current_section:
                pages.append(ExtractedPage(
                    page_number=page_number,
                    content=current_section,
                    content_type="text",
                    metadata={"level": current_level, "section": current_section, "is_heading": True}
                ))
        else:
            text = para.text.strip()
            if text:
                pages.append(ExtractedPage(
                    page_number=page_number,
                    content=text,
                    content_type="text",
                    metadata={"level": current_level, "section": current_section}
                ))
                
    for table in doc.tables:
        markdown_table = []
        for i, row in enumerate(table.rows):
            row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            markdown_table.append("| " + " | ".join(row_data) + " |")
            if i == 0:
                markdown_table.append("|" + "|".join(["---"] * len(row_data)) + "|")
        
        if markdown_table:
            pages.append(ExtractedPage(
                page_number=page_number,
                content="\n".join(markdown_table),
                content_type="table",
                metadata={"level": current_level, "section": current_section}
            ))
            
    return pages
